"""Pruebas unitarias de formatos documentales sin red ni Gemini."""

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from Agente.app.procesamiento import (
    DocumentoDescubierto,
    ErrorExtraccionDocumento,
    extraer_documento,
    fragmentar_documento,
)
from Agente.app.procesamiento.extractores import obtener_extractor


class ExtractoresDocumentosTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporal = tempfile.TemporaryDirectory()
        self.raiz = Path(self.temporal.name)

    def tearDown(self) -> None:
        self.temporal.cleanup()

    def _documento(self, nombre: str) -> DocumentoDescubierto:
        return DocumentoDescubierto(
            empresa="EmpresaPrueba",
            visibilidad="Public",
            ruta_relativa=f"EmpresaPrueba/Public/{nombre}",
            ruta_archivo=self.raiz / nombre,
        )

    def test_selector_elige_extractor_por_extension(self) -> None:
        esperados = {
            "manual.md": "extraer_markdown",
            "manual.markdown": "extraer_markdown",
            "notas.txt": "extraer_txt",
            "manual.docx": "extraer_docx",
            "manual.pdf": "extraer_pdf",
        }
        for nombre, funcion in esperados.items():
            with self.subTest(nombre=nombre):
                self.assertEqual(obtener_extractor(nombre).__name__, funcion)

    def test_selector_rechaza_doc_y_extension_desconocida(self) -> None:
        with self.assertRaisesRegex(
            ErrorExtraccionDocumento,
            "Convierte el archivo a .docx",
        ):
            obtener_extractor("legado.doc")
        with self.assertRaisesRegex(ErrorExtraccionDocumento, "no soportado"):
            obtener_extractor("datos.csv")

    def test_txt_se_normaliza_al_modelo_interno(self) -> None:
        documento = self._documento("notas.txt")
        documento.ruta_archivo.write_text(
            "Primera línea\r\n\r\n\r\nSegunda línea\r\n",
            encoding="utf-8",
            newline="",
        )

        extraido = extraer_documento(documento)

        self.assertEqual(extraido.contenido_markdown, "Primera línea\n\nSegunda línea")
        self.assertEqual(extraido.tipo_archivo, "txt")
        self.assertEqual(extraido.archivo_original, "notas.txt")
        self.assertEqual(len(extraido.secciones), 1)

    def test_docx_conserva_encabezados_listas_y_tablas(self) -> None:
        from docx import Document

        documento = self._documento("manual.docx")
        archivo = Document()
        archivo.add_heading("Servicios", level=1)
        archivo.add_paragraph("Descripción general")
        archivo.add_paragraph("Consultoría", style="List Bullet")
        tabla = archivo.add_table(rows=2, cols=2)
        tabla.cell(0, 0).text = "Servicio"
        tabla.cell(0, 1).text = "Estado"
        tabla.cell(1, 0).text = "RAG"
        tabla.cell(1, 1).text = "Activo"
        archivo.save(documento.ruta_archivo)

        extraido = extraer_documento(documento)

        self.assertEqual(extraido.tipo_archivo, "docx")
        self.assertEqual(extraido.secciones[0].titulo, "Servicios")
        self.assertIn("# Servicios", extraido.contenido_markdown)
        self.assertIn("- Consultoría", extraido.contenido_markdown)
        self.assertIn("| Servicio | Estado |", extraido.contenido_markdown)
        self.assertIn("| RAG | Activo |", extraido.contenido_markdown)

    @patch("pypdf.PdfReader")
    def test_pdf_nativo_conserva_paginas_y_metadatos(self, lector_mock) -> None:
        documento = self._documento("politicas.pdf")
        documento.ruta_archivo.write_bytes(b"%PDF-contenido-ficticio")
        lector_mock.return_value = SimpleNamespace(
            is_encrypted=False,
            pages=[
                SimpleNamespace(extract_text=lambda: "Política general"),
                SimpleNamespace(extract_text=lambda: "Segunda página"),
            ],
        )

        extraido = extraer_documento(documento)
        fragmentos = fragmentar_documento(extraido)

        self.assertEqual([seccion.pagina for seccion in extraido.secciones], [1, 2])
        self.assertEqual(extraido.secciones[0].titulo, "Página 1")
        self.assertEqual(fragmentos[0].pagina, 1)
        self.assertEqual(fragmentos[0].tipo_archivo, "pdf")
        self.assertEqual(fragmentos[0].archivo_original, "politicas.pdf")
        self.assertEqual(fragmentos[0].empresa, "EmpresaPrueba")
        self.assertEqual(fragmentos[0].visibilidad, "Public")
        self.assertEqual(fragmentos[0].ruta_relativa, documento.ruta_relativa)
        self.assertIn("#seccion-1-fragmento-1", fragmentos[0].referencia_fragmento)

    @patch("pypdf.PdfReader")
    def test_pdf_sin_texto_se_rechaza_con_mensaje_claro(self, lector_mock) -> None:
        documento = self._documento("escaneado.pdf")
        documento.ruta_archivo.write_bytes(b"%PDF-imagen-ficticia")
        lector_mock.return_value = SimpleNamespace(
            is_encrypted=False,
            pages=[SimpleNamespace(extract_text=lambda: "   ")],
        )

        with self.assertRaisesRegex(
            ErrorExtraccionDocumento,
            "PDF escaneado.*OCR aún no está soportado",
        ):
            extraer_documento(documento)


if __name__ == "__main__":
    unittest.main()
